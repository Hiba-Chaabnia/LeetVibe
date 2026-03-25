"""DOCX export for the LeetVibe Playbook Reference Guide.

Requires python-docx (listed in project dependencies).
Call export_reference_docx(topics, notes) → returns the output file path string.
"""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path


def _desktop() -> Path:
    """Return the user's Desktop folder, resolving OneDrive redirects on Windows."""
    if sys.platform == "win32":
        try:
            import winreg
            key = winreg.OpenKey(
                winreg.HKEY_CURRENT_USER,
                r"Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders",
            )
            desktop, _ = winreg.QueryValueEx(key, "Desktop")
            winreg.CloseKey(key)
            p = Path(desktop)
            if p.exists():
                return p
        except Exception:
            pass
    # Fallback chain for non-Windows or registry failure
    for candidate in [
        Path.home() / "Desktop",
        Path.home() / "OneDrive" / "Desktop",
        Path.home() / "OneDrive" / "Bureau",
        Path.home(),
    ]:
        if candidate.exists():
            return candidate
    return Path.home()


# ── Brand colours (RGB tuples) ─────────────────────────────────────────────────
_FIRE       = (0xFF, 0x82, 0x05)   # primary orange
_GOLD       = (0xFF, 0xD7, 0x00)   # section headers / notes label
_AMBER      = (0xFF, 0xB3, 0x00)   # complexity values
_DIM_CODE   = (0xCC, 0xCC, 0xCC)   # code text colour
_CODE_FILL  = "1E1E1E"             # code block background (hex, no #)
_PAGE_BG    = "121212"             # page background


def export_reference_docx(topics: list[dict], notes: dict[str, str]) -> str:
    """Export *topics* + *notes* to a formatted DOCX file, grouped by category.

    Saves to ~/Desktop/leetvibe_reference.docx and returns that path as str.
    Raises ImportError if python-docx is not installed.
    """
    from docx import Document                         # type: ignore[import]
    from docx.enum.text import WD_ALIGN_PARAGRAPH     # type: ignore[import]
    from docx.oxml import OxmlElement                 # type: ignore[import]
    from docx.oxml.ns import qn                       # type: ignore[import]
    from docx.shared import Inches, Pt, RGBColor      # type: ignore[import]

    from .data.topics import CATEGORIES

    _DIFF_FULL = {"E": "Easy", "M": "Medium", "H": "Hard"}

    # Build lookup: title → topic dict
    topic_by_title = {t["title"]: t for t in topics}

    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Cover page ──────────────────────────────────────────────────────
    _cover(doc, WD_ALIGN_PARAGRAPH, Pt, RGBColor)

    first_block = True

    # ── Topics grouped by category ───────────────────────────────────────
    for cat in CATEGORIES:
        cat_topics = [topic_by_title[t] for t in cat["topics"] if t in topic_by_title]
        if not cat_topics:
            continue

        if not first_block:
            doc.add_page_break()
        first_block = False

        # Category heading
        cat_heading = doc.add_heading(f"{cat['icon']}  {cat['name']}", level=1)
        _color_runs(cat_heading, RGBColor(*_FIRE))
        _set_para_spacing(cat_heading, before=0, after=8)
        _hr(doc, OxmlElement, qn, RGBColor(*_FIRE))

        for topic in cat_topics:
            note = notes.get(topic["slug"], "").strip()

            # Topic title (H2)
            h = doc.add_heading(topic["title"], level=2)
            _color_runs(h, RGBColor(*_GOLD))
            _set_para_spacing(h, before=12, after=4)

            # Recognised by
            recognize = topic.get("recognize", "").replace("\n  ", " ").strip()
            if recognize:
                _section_label(doc, "Recognised by", Pt, RGBColor(*_GOLD))
                _body(doc, recognize, Pt)

            # Diagram
            if topic.get("diagram"):
                _section_label(doc, "Diagram", Pt, RGBColor(*_GOLD))
                _code_block(doc, topic["diagram"], Pt, RGBColor, OxmlElement, qn)

            # When to use
            if topic.get("when"):
                _section_label(doc, "When to use", Pt, RGBColor(*_GOLD))
                _body(doc, topic["when"].replace("\n  ", " "), Pt)

            # Patterns (numbered list)
            patterns = topic.get("patterns", [])
            if patterns:
                _section_label(doc, "Patterns", Pt, RGBColor(*_GOLD))
                for i, pat in enumerate(patterns, 1):
                    np = doc.add_paragraph()
                    r = np.add_run(f"{i}. {pat['name']}")
                    r.bold = True
                    r.font.size = Pt(10)
                    _set_para_spacing(np, before=4, after=2)
                    _code_block(doc, pat["code"], Pt, RGBColor, OxmlElement, qn)

            # Complexity
            t_val = topic.get("time", "").strip()
            s_val = topic.get("space", "").strip()
            if t_val or s_val:
                _section_label(doc, "Complexity", Pt, RGBColor(*_GOLD))
                p = doc.add_paragraph()
                _set_para_spacing(p, before=0, after=2)
                if t_val:
                    _run(p, "Time:   ", Pt(10), bold=True)
                    _run(p, t_val + "    ", Pt(10), color=RGBColor(*_AMBER))
                if s_val:
                    _run(p, "Space:  ", Pt(10), bold=True)
                    _run(p, s_val, Pt(10), color=RGBColor(*_AMBER))

            # Pitfalls
            pitfalls = topic.get("pitfalls", "").strip()
            if pitfalls:
                _section_label(doc, "Pitfalls", Pt, RGBColor(*_GOLD))
                _body(doc, pitfalls, Pt)

            # Classic problems
            if topic.get("problems"):
                _section_label(doc, "Classic Problems", Pt, RGBColor(*_GOLD))
                for prob in topic["problems"]:
                    if isinstance(prob, tuple):
                        name, diff = prob
                        label = f"{name}  [{_DIFF_FULL.get(diff, diff)}]"
                    else:
                        label = str(prob)
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.add_run(label).font.size = Pt(10)
                    _set_para_spacing(bp, before=0, after=1)

            # Notes
            if note:
                _section_label(doc, "My Notes", Pt, RGBColor(*_GOLD))
                np = doc.add_paragraph()
                np.add_run(note).font.size = Pt(10)
                _set_para_spacing(np, before=0, after=4)

    # ── Save ────────────────────────────────────────────────────────────
    out = _desktop() / "leetvibe_reference.docx"
    doc.save(str(out))
    return str(out)


# ── Private helpers ────────────────────────────────────────────────────────────

def _cover(doc, WD_ALIGN, Pt, RGBColor) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN.CENTER
    r = p.add_run("LeetVibe")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(*_FIRE)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN.CENTER
    r2 = p2.add_run("Playbook Reference Guide")
    r2.bold = True
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(*_GOLD)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN.CENTER
    r3 = p3.add_run(f"Generated  {date.today().isoformat()}")
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()   # spacer
    doc.add_page_break()


def _hr(doc, OxmlElement, qn, color) -> None:
    """Draw a thin coloured horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "FF8205")
    pBdr.append(bottom)
    pPr.append(pBdr)
    _set_para_spacing(p, before=0, after=4)


def _section_label(doc, text: str, Pt, color) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = color
    _set_para_spacing(p, before=6, after=2)


def _body(doc, text: str, Pt) -> None:
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10)
    _set_para_spacing(p, before=0, after=4)


def _code_block(doc, code: str, Pt, RGBColor, OxmlElement, qn) -> None:
    """Render a multi-line code string with dark background and Courier font."""
    for line in code.split("\n"):
        p = doc.add_paragraph()
        r = p.add_run(line if line.strip() else " ")
        r.font.name = "Courier New"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(*_DIM_CODE)
        # Dark background shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  _CODE_FILL)
        pPr.append(shd)
        _set_para_spacing(p, before=0, after=0)
    # Spacer after block
    sp = doc.add_paragraph()
    _set_para_spacing(sp, before=0, after=4)


def _run(p, text: str, size, bold: bool = False, color=None) -> None:
    r = p.add_run(text)
    r.font.size = size
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color


def _color_runs(para, color) -> None:
    for r in para.runs:
        r.font.color.rgb = color


def _set_para_spacing(para, before: int = 0, after: int = 4) -> None:
    from docx.shared import Pt  # type: ignore[import]
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
